# src/utils.py
"""Utility functions for file generation and other helpers."""

import streamlit as st
from io import BytesIO

# --- Document Generation for Download ---

def generate_docx_bytes(story_content: str) -> bytes:
    """Generates a .docx file content in memory."""
    try:
        from docx import Document
    except ImportError:
        st.error("python-docx library not installed. Cannot generate Word document.")
        return b"" # Return empty bytes if library is missing

    try:
        document = Document()
        document.add_heading('Generated Story', level=1)
        document.add_paragraph(story_content)

        bio = BytesIO()
        document.save(bio)
        bio.seek(0)
        return bio.getvalue()
    except Exception as e:
        st.error(f"Error generating Word document: {e}")
        return b""

def generate_pdf_bytes_from_text(story_content: str) -> bytes:
    """Generates a simple PDF file content in memory from text."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
    except ImportError:
        st.error("reportlab library not installed. Cannot generate PDF document.")
        return b"" # Return empty bytes if library is missing

    try:
        bio = BytesIO()
        doc = SimpleDocTemplate(bio, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Add title
        title_style = styles['h1']
        story.append(Paragraph("Generated Story", title_style))

        # Add content, handle newlines
        body_style = styles['BodyText']
        paragraphs = story_content.split('\n')
        for para in paragraphs:
            story.append(Paragraph(para, body_style))
            story.append(Paragraph("<br/>", body_style)) # Add space between paragraphs

        doc.build(story)
        bio.seek(0)
        return bio.getvalue()
    except Exception as e:
        st.error(f"Error generating PDF document: {e}")
        return b""

# --- Other potential utils ---

def clean_text(text: str) -> str:
    """Basic text cleaning function (example)."""
    # Add any text cleaning logic needed before sending to AI
    import re
    text = re.sub(r'\s+', ' ', text).strip() # Replace multiple spaces with one
    return text

